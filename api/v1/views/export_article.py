"""Export Article Logic"""
import io
from flask import Flask, Blueprint, send_file, request
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

export = Blueprint('export', __name__)

@export.route('/export', methods=['POST'], strict_slashes=False)
def export_content():
    """Export the content to desired format"""
    content = request.form.get('content')
    topic = request.form.get('topic')
    export_format = request.form.get('exportFormat')

    if export_format == 'docx':
        file_data = export_docx(content, topic)
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif export_format == 'pdf':
        file_data = export_pdf(content, topic)
        mime_type = "application/pdf"
    elif export_format == 'md':
        file_data = export_markdown(content, topic)
        mime_type = "text/markdown"
    elif export_format == 'txt':
        file_data = export_txt(content)
        mime_type = "text/plain"
    else:
        return "Unsupported Format", 400

    return send_file(file_data,
                     as_attachment=True,
                     download_name=f"{topic}.{export_format}",
                     mimetype=mime_type
                     )

def export_docx(content, topic):
    """Export content as DOCX"""
    doc = Document()
    doc.add_heading(topic, level=1)
    doc.add_paragraph(content)
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def export_pdf(content, topic):
    """Export content as PDF"""
    file_stream = io.BytesIO()
    pdf = SimpleDocTemplate(file_stream, pagesize=letter) # Create a SimpleDocTemplate object
    styles = getSampleStyleSheet() # Get standard PDF styles
    elements = [] # Create a list to hold the elements to be added to the PDF
    elements.append(Paragraph(topic, styles['Title'])) # Add the topic as a heading to the elements list
    elements.append(Spacer(1, 12)) # Add some space after the heading
    
    # Split the content into paragraphs and add each to the elements list
    for paragraph in content.split('\n\n'):  # Assuming new paragraphs are separated by two newlines
        elements.append(Paragraph(paragraph, styles['Normal']))
        elements.append(Spacer(1, 12))

    pdf.build(elements) # Build the PDF with the list of elements
    file_stream.seek(0) # Move the file pointer to the beginning
    return file_stream # Return the file stream

def export_markdown(content, topic):
    """Export content as Markdown"""
    markdown_content = (f"#{topic}\n\n{content}")
    file_stream = io.BytesIO(markdown_content.encode('utf-8')) # Encode string to bytes!!
    return file_stream

def export_txt(content):
    """Export content as Txt"""
    file_stream = io.BytesIO(content.encode('utf-8')) # Encode string to bytes!!
    return file_stream
